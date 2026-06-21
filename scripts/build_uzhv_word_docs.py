# -*- coding: utf-8 -*-
"""Generate formatted Word documents for AIS UZHV planning materials."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "word"

COLOR_PRIMARY = RGBColor(0x1A, 0x47, 0x7A)  # navy
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
COLOR_MUTED = RGBColor(0x5A, 0x5A, 0x5A)
COLOR_CALLOUT_BG = "E8F1FA"
COLOR_TABLE_HEADER = "1A477A"
COLOR_TABLE_ALT = "F2F6FA"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = COLOR_PRIMARY


def add_footer(doc: Document, text: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED


def add_title_page(doc: Document, title: str, subtitle: str, meta_lines: list[str]) -> None:
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = COLOR_PRIMARY
    r.font.name = "Calibri"

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(subtitle)
    rs.font.size = Pt(14)
    rs.font.color.rgb = COLOR_ACCENT
    rs.font.name = "Calibri"

    doc.add_paragraph()
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(meta_lines):
        if i:
            box.add_run("\n")
        run = box.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_MUTED

    doc.add_page_break()


def add_callout(doc: Document, text: str, bold_parts: bool = True) -> None:
    p = doc.add_paragraph()
    set_paragraph_shading(p, COLOR_CALLOUT_BG)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if bold_parts:
        run.bold = False


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, COLOR_TABLE_HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        if ri % 2 == 1:
            for c in cells:
                set_cell_shading(c, COLOR_TABLE_ALT)
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


def strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", s)


def build_plan() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title_page(
        doc,
        "Программа реализации\nАИС УЖВ",
        "Для представления заказчику · май — октябрь 2026",
        [
            "Документ-основание: ТЗ_АИС_УЖВ (2).docx",
            "Заказчик: Управление по жилищным вопросам",
            "администрации МО город Краснодар",
            "Оператор закупки: МКУ «Электронный Краснодар»",
            "Версия 2.0 · 19.05.2026",
        ],
    )

    add_callout(
        doc,
        "Важно: в упрощённом ТЗ внешние интеграции (ГИС ЖКХ, СМЭВ, Катарсис и др.) "
        "не входят в объём первого этапа — опции I-01…I-14 не активны. "
        "Учёт межведомственных ответов ведётся вручную; обмен — файлами CSV/XLSX и скриптами PY-01…PY-07.\n\n"
        "Срок по ТЗ: оказание услуг до 02.10.2026. Приёмка планируется к концу сентября — началу октября. "
        "Перенос на ноябрь возможен только изменением ТЗ/контракта.",
    )

    add_h(doc, "1. Чем упрощённое ТЗ отличается от полной версии 1.1", 1)
    add_table(
        doc,
        ["Аспект", "ТЗ v1.1", "Согласованное ТЗ (2)"],
        [
            ["Интеграции", "ГИС ЖКХ, СМЭВ, Катарсис", "Нет — опции I-xx не активны"],
            ["Межведомственные запросы", "Авто через СМЭВ", "Ручной учёт + отчёт ОТЧ-8"],
            ["Обмен данными", "API внешних ИС", "Импорт/экспорт файлов, Python"],
            ["Срок контракта", "В договоре", "До 02.10.2026"],
            ["Модули учёта, договоров, контроля", "Да", "Да (функционал сохранён)"],
        ],
        [4.5, 5.5, 6.5],
    )
    add_p(doc, "Для заказчика это плюс: реалистичнее уложиться в календарь, меньше зависимость от смежных ведомств.")

    add_h(doc, "2. Сводный календарь", 1)
    add_table(
        doc,
        ["Период", "Работы", "Результат"],
        [
            ["Май — июнь", "Прототип, 4 визита в УЖВ", "Демо + протоколы"],
            ["Конец июня (1 нед.)", "Регистрация прав на ПО", "Заявка в Роспатент"],
            ["Июль", "Реестр Минцифры + разработка", "Запись в реестре"],
            ["Август", "Реестровый № в ТЗ, закупка", "Извещение в ЕИС"],
            ["Сентябрь", "Торги, контракт, внедрение", "Акт начала работ"],
            ["Октябрь (до 02.10)", "Опытная эксплуатация, приёмка", "Акт приёма-передачи"],
            ["Ноябрь", "Только при изменении срока", "Резерв / гарантия"],
        ],
        [3.2, 6.5, 5.0],
    )

    add_h(doc, "3. Май — июнь: прототип и 4 очные сессии", 1)
    add_h(doc, "3.1. Состав прототипа (к концу июня)", 2)
    add_p(doc, "Платформа: веб, JWT, роли, журнал аудита, PostgreSQL, Docker, Astra Linux, Nginx.")
    add_table(
        doc,
        ["Подсистема", "В прототипе (июнь)", "К приёмке 02.10"],
        [
            ["Малоимущие + расчёт", "Заявление, расчёт, заключение", "Полный цикл"],
            ["Учёт нуждающихся", "Реестры, очерёдность", "+ история, отчёты"],
            ["Молодые семьи", "Списки (базово)", "По законам КК 2704, 2710"],
            ["Дети-сироты", "Карточка дела, ручные решения", "Без I-04, I-08"],
            ["Договоры", "Реестр договоров", "+ шаблоны docx/pdf"],
            ["Жилфонд", "МКД, помещения", "Импорт PY-01"],
            ["Жилконтроль", "Проверки, предписания", "Реестр проверок"],
            ["Обращения", "Регистрация, сроки 30 дн.", "+ ответы, ОТЧ-5"],
            ["Отчёты", "ОТЧ-1, ОТЧ-5", "ОТЧ-1…ОТЧ-10"],
            ["Интеграции I-xx", "Не делаем", "Заглушки / справочник"],
        ],
        [4.0, 5.5, 5.5],
    )

    add_h(doc, "3.2. График визитов", 2)
    add_table(
        doc,
        ["№", "Окно", "Участники", "Цель"],
        [
            ["1", "26–30.05", "Руководство, ИТ", "Обследование, приоритеты"],
            ["2", "09–13.06", "Учёт, малоимущие", "Поля, расчёты, формы"],
            ["3", "16–20.06", "Ключевые пользователи", "Утверждение UI"],
            ["4", "23–27.06", "УЖВ + закупки", "Итог MVP, реестр, август"],
        ],
        [1.0, 2.5, 4.5, 6.5],
    )

    add_h(doc, "3.3. Блоки вопросов для анкетирования (14+)", 2)
    questions = [
        "Актуальные регламенты и приказы (в т.ч. пр. № 58/2026).",
        "Образцы заявлений, заключений, постановлений.",
        "Алгоритм расчёта дохода и имущества.",
        "Категории учёта и очерёдность.",
        "Формы ОТЧ-1…ОТЧ-10 — обязательные к 02.10.",
        "Роли и матрица утверждений.",
        "Регламент обращений (30 дней).",
        "Дела детей-сирот без СМЭВ — как фиксируют ответы ведомств.",
        "Жилфонд: УЖВ / Горжилхоз, формат выгрузок.",
        "Жилищный контроль — предметы проверок.",
        "Контур размещения (сервер, VPN).",
        "Требования ИБ и 152-ФЗ.",
        "Способ закупки и НМЦК.",
        "Обезличенные данные для пилота (PY-04).",
    ]
    add_bullets(doc, questions)

    add_h(doc, "4. Конец июня: регистрация прав (1 неделя)", 1)
    add_table(
        doc,
        ["День", "Действие"],
        [
            ["1–2", "Заявка «программа для ЭВМ», реферат, фрагмент кода"],
            ["3", "Подача в Роспатент"],
            ["4–7", "Номер заявки → основание для реестра"],
        ],
        [2.5, 12.0],
    )

    add_h(doc, "5. Июль: реестр + разработка", 1)
    add_bullets(
        doc,
        [
            "Подача в реестр → запись до конца июля / начала августа.",
            "Все подсистемы — рабочие сценарии.",
            "PY-01, PY-03, PY-04 — импорт и миграция.",
            "PY-02, PY-06 — отчёты и пакеты документов.",
            "Шаблонизатор docx/pdf, CI/CD, Docker, документация.",
        ],
    )

    add_h(doc, "6. Август: реестровый номер и закупка", 1)
    add_bullets(
        doc,
        [
            "Получение реестровой записи.",
            "Дополнение описания закупки: функции, реестр, без I-xx, протоколы визитов.",
            "Публикация закупки в ЕИС.",
            "Обоснование НМЦК.",
        ],
    )

    add_h(doc, "7. Сентябрь: торги и внедрение", 1)
    add_table(
        doc,
        ["Неделя", "Событие"],
        [
            ["1–2", "Завершение процедуры, контракт"],
            ["3", "Развёртывание (Docker, PostgreSQL, Nginx)"],
            ["4", "Миграция PY-04, обучение"],
        ],
        [2.5, 12.0],
    )

    add_h(doc, "8. Октябрь: приёмка (до 02.10.2026)", 1)
    add_table(
        doc,
        ["Вид испытаний", "Содержание"],
        [
            ["Функциональные", "Все подсистемы без I-xx"],
            ["Нагрузочные", "50 пользователей, отклик ≤ 3 с"],
            ["ИБ", "JWT, TLS, журнал, 152-ФЗ"],
            ["Интеграции", "Не тестируются"],
            ["Опытная эксплуатация", "≥ 14 дней"],
            ["Приёмочные", "Акт приёма-передачи"],
        ],
        [5.0, 10.5],
    )

    add_h(doc, "9. Этап 2 — опциональные интеграции", 1)
    add_table(
        doc,
        ["Код", "Система", "Когда"],
        [
            ["I-01", "ГИС ЖКХ", "По готовности тех. условий"],
            ["I-08", "Катарсис / СМЭВ", "После согласования с Минтрудом КК"],
            ["I-05, I-06", "Госуслуги / ЕСИА", "По инфраструктуре"],
        ],
        [2.0, 5.5, 6.0],
    )

    add_h(doc, "10. Риски", 1)
    add_table(
        doc,
        ["Риск", "Митигация"],
        [
            ["Ожидание интеграций как в полном ТЗ", "Письменно: ТЗ (2) без I-xx"],
            ["Срок 02.10 vs ноябрь", "Согласовать в контракте"],
            ["Большой объём модулей", "Приоритизация на визите 4, PY"],
            ["Задержка реестра", "Подача в начале июля"],
        ],
        [6.5, 8.0],
    )

    add_h(doc, "11. Тезис для руководства УЖВ", 1)
    add_p(
        doc,
        "По согласованному упрощённому ТЗ создаётся АИС УЖВ с полным внутренним функционалом "
        "без подключения внешних систем на первом этапе. С мая по июнь — прототип и четыре сессии "
        "согласования; в июле — реестр; в августе — закупка; в сентябре–октябре — внедрение и "
        "приёмка до 02.10.2026. Интеграции — этап 2.",
        italic=True,
    )

    add_footer(doc, "АИС УЖВ · Программа реализации v2.0 · 19.05.2026")
    return doc


def build_visits() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title_page(
        doc,
        "Протоколы очных сессий",
        "АИС УЖВ · шаблоны для визитов в УЖВ",
        ["Май — июнь 2026", "4 рабочие сессии", "Версия 1.0 · 19.05.2026"],
    )

    visits = [
        (
            "Визит 1 — Обследование (26–30.05.2026)",
            "Процессы «как есть», системы, приоритеты этапа 1.",
            ["Презентация целей АИС УЖВ и календаря.", "Обход процессов: заявление → учёт → отчёт.", "Демонстрация существующих систем.", "Фиксация болей и KPI."],
            ["Регламенты УЖВ — актуальные редакции", "Образцы заявлений и заключений", "Выгрузка из текущей ИС (обезличенная)", "Схема сети / требования ИТ", "Контакт по закупке и ИБ"],
            [["1", "Приоритет модулей на этап 1", ""], ["2", "Контур размещения", ""], ["3", "Срок следующего визита", ""]],
        ),
        (
            "Визит 2 — Модули учёта (09–13.06.2026)",
            "Поля данных, расчёты, печатные формы.",
            ["Сверка карточки заявителя и семьи.", "Алгоритм расчёта дохода (2–3 кейса).", "Алгоритм имущества.", "Очерёдность, снятие с учёта.", "Перечень отчётов этапа 1."],
            [],
            [["Кейс А", "Семья 3 чел., доход ниже прожиточного", "Признание малоимущими"], ["Кейс Б", "Транспорт выше нормы", "Отказ / особое заключение"], ["Кейс В", "Изменение состава семьи", "Обновление учётного дела"]],
        ),
        (
            "Визит 3 — Утверждение интерфейсов (16–20.06.2026)",
            "Подписать согласованные экраны прототипа.",
            ["Рабочий стол специалиста", "Реестр заявлений", "Карточка заявления", "Расчёт дохода/имущества", "Учётное дело", "Обращение гражданина", "Отчёты", "Администрирование"],
            [],
            [],
        ),
        (
            "Визит 4 — Итог MVP и закупка (23–27.06.2026)",
            "Закрыть MVP, согласовать этапность и график закупки.",
            ["Демонстрация прототипа по замечаниям.", "Объём этапа 1 vs этап 2.", "Реестр ПО и сроки.", "Вопросы закупочного подразделения.", "Утверждение календарного плана."],
            [],
            [["1", "Утвердить объём этапа 1", ""], ["2", "Комиссия по приёмке", ""], ["3", "Данные для НМЦК", ""], ["4", "Дата публикации закупки", "август 2026"]],
        ),
    ]

    for title, goal, agenda, checklist, table_rows in visits:
        add_h(doc, title, 1)
        add_p(doc, f"Цель: {goal}")
        add_h(doc, "Повестка", 2)
        add_bullets(doc, agenda)
        if checklist:
            add_h(doc, "Чек-лист материалов", 2)
            add_bullets(doc, checklist)
        if table_rows:
            if "Кейс" in table_rows[0][0]:
                add_h(doc, "Кейсы для прогона", 2)
                add_table(doc, ["Кейс", "Вход", "Ожидаемый результат"], table_rows, [2.5, 5.5, 6.5])
            elif table_rows[0][0] == "1" and "Утвердить" in str(table_rows):
                add_h(doc, "Итоговые решения", 2)
                add_table(doc, ["№", "Решение", "Срок"], table_rows, [1.0, 8.0, 4.0])
            else:
                add_h(doc, "Решения (заполнить на месте)", 2)
                add_table(doc, ["№", "Вопрос", "Решение"], table_rows, [1.0, 6.0, 6.5])

        if "Визит 3" in title:
            add_h(doc, "Протокол утверждения интерфейса", 2)
            add_callout(
                doc,
                "Интерфейсы перечисленных экранов прототипа АИС УЖВ от «___» _________ 2026 г. "
                "согласованы для реализации в этапе 1 с замечаниями из приложения № ___.\n\n"
                "Подписи: _________________ (УЖВ)    _________________ (разработчик)",
            )
        doc.add_page_break()

    add_footer(doc, "АИС УЖВ · Шаблоны протоколов v1.0")
    return doc


def build_calendar() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title_page(
        doc,
        "Календарный план\nАИС УЖВ",
        "Май — октябрь 2026 · рабочая версия",
        ["Упрощённое ТЗ без активных интеграций", "Приёмка по ТЗ: до 02.10.2026", "19.05.2026"],
    )

    add_callout(doc, "Основание: ТЗ_АИС_УЖВ (2).docx. Интеграции I-01…I-14 не активны. Не юридическая консультация.")

    add_h(doc, "Сравнение вариантов календаря", 1)
    add_table(
        doc,
        ["Месяц", "Исходный вариант", "Скорректировано"],
        [
            ["Июнь", "Попадание в реестр", "Подача в реестр; запись — июль"],
            ["Июль", "Только прототип", "Запись в реестре + доработка"],
            ["Август", "Публикация закупки", "Извещение в ЕИС"],
            ["Сен–Окт", "Доработка до контракта", "Работы по контракту"],
            ["Ноябрь", "Приёмка всего ТЗ", "Приёмка этапа 1 / по договору"],
        ],
        [2.5, 5.5, 6.5],
    )

    months = [
        ("Май 2026", ["Разбор ТЗ, архитектура", "MVP: обращения, роли, статусы, отчёты", "Демо для УЖВ"], "URL демо + протокол"),
        ("Июнь 2026", ["Доработка MVP", "Комплект на реестр", "Подача в реестр Минцифры"], "Заявка подана; MVP согласован"),
        ("Июль 2026", ["Ответы на замечания реестра", "Запись в реестре", "Финал ТЗ и НМЦК"], "№ записи реестра"),
        ("Август 2026", ["Публикация в ЕИС", "Подготовка заявки участника"], "Извещение опубликовано"),
        ("Сентябрь 2026", ["Итоги торгов, контракт", "Развёртывание в контуре УЖВ"], "Контракт подписан"),
        ("Октябрь 2026", ["Доработка этапа 1", "Обучение", "Опытная эксплуатация"], "Акт / промежуточная приёмка"),
    ]
    add_h(doc, "Помесячный график", 1)
    for name, tasks, checkpoint in months:
        add_h(doc, name, 2)
        add_bullets(doc, tasks)
        add_p(doc, f"Контрольная точка: {checkpoint}.")

    add_h(doc, "Объём этапа 1", 1)
    add_table(
        doc,
        ["Входит в этап 1", "Не входит (этап 2+)"],
        [
            ["Лицензии, развёртывание", "ГИС ЖКХ, СМЭВ"],
            ["Обращения, роли, статусы", "Полная мобильная витрина"],
            ["Базовые отчёты (2–5)", "Сложная аналитика / BI"],
            ["Обучение, опытная эксплуатация", "Полная миграция архива"],
        ],
        [7.0, 7.5],
    )

    add_h(doc, "Риски по срокам", 1)
    add_table(
        doc,
        ["Риск", "Вероятность", "Митигация"],
        [
            ["Реестр сдвигается", "Высокая", "Подача в июне"],
            ["Торги в сентябре", "Средняя", "Извещение в начале августа"],
            ["Проигрыш торгов", "Средняя", "План B: субподряд"],
            ["Теневая разработка", "Юридический", "Только демо до контракта"],
        ],
        [5.5, 3.0, 6.0],
    )

    add_footer(doc, "АИС УЖВ · Календарный план · 19.05.2026")
    return doc


def parse_markdown_to_doc(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if len(table_buf) >= 2:
            headers = [c.strip() for c in table_buf[0].strip("|").split("|")]
            rows = []
            for row in table_buf[2:]:
                rows.append([c.strip() for c in row.strip("|").split("|")])
            if headers and rows:
                add_table(doc, headers, rows)
        table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("|") and "|" in line[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", line):
                i += 1
                continue
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if line.startswith("# ") and not line.startswith("## "):
            if "Часть " in line:
                doc.add_page_break()
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("> "):
            add_callout(doc, line[2:].strip(), bold_parts=False)
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ " + line[6:].strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(strip_md_bold(line[2:].strip()))
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(strip_md_bold(re.sub(r"^\d+\.\s", "", line).strip()))
        elif line.strip() == "---":
            pass
        elif line.strip():
            p = doc.add_paragraph()
            p.add_run(strip_md_bold(line.strip()))
        i += 1
    flush_table()


def build_questionnaire() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title_page(
        doc,
        "Опросник для заказчика\nАИС УЖВ",
        "УЖВ · ИКТ · коммерция · май–июнь 2026",
        [
            "УЖВ — функционал",
            "ИКТ — закупка, сервер, приёмка",
            "Платформа в реестре + модуль УЖВ",
            "Версия 1.1 · 19.05.2026",
        ],
    )
    md_path = ROOT / "docs" / "voprosy-k-zakazchiku-uzhv-detalno.md"
    parse_markdown_to_doc(doc, md_path.read_text(encoding="utf-8"))
    add_footer(doc, "АИС УЖВ · Опросник для заказчика v1.0")
    return doc


def build_from_md(md_name: str, title: str, subtitle: str) -> Document:
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, title, subtitle, ["Версия 1.0 · 19.05.2026"])
    md_path = ROOT / "docs" / md_name
    parse_markdown_to_doc(doc, md_path.read_text(encoding="utf-8"))
    add_footer(doc, f"АИС УЖВ · {md_name}")
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    files = [
        (
            "delayu-product-opisanie.docx",
            lambda: build_from_md(
                "delayu-product-opisanie.md",
                "ДелаЮ\n(Дела.ЮГИт)",
                "Описание продукта · ЮГИт",
            ),
        ),
        (
            "goszakazchik-komplekt-2026.docx",
            lambda: build_from_md(
                "goszakazchik-komplekt-2026.md",
                "Комплект для\nгосзаказчика",
                "УЖВ · ИКТ · платформа · май–октябрь 2026",
            ),
        ),
        (
            "kommerciya-komplekt-2026.docx",
            lambda: build_from_md(
                "kommerciya-komplekt-2026.md",
                "Комплект для\nкоммерции",
                "Пилоты · Ателика · стройка · юрбюро",
            ),
        ),
        ("plan-dlya-zakazchika-ais-uzhv-2026.docx", build_plan),
        ("vizity-uzhv-protokoly-shablon.docx", build_visits),
        ("uzhv-calendar-2026.docx", build_calendar),
        ("voprosy-k-zakazchiku-uzhv-detalno.docx", build_questionnaire),
    ]
    for name, builder in files:
        path = OUT / name
        builder().save(path)
        print(f"Saved: {path}")


if __name__ == "__main__":
    main()

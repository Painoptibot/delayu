"""Печатные формы УЖВ (заключения, проекты постановлений)."""
from __future__ import annotations

import io
import re
from typing import Any

from delayu.models import PrintTemplate
from delayu.models_uzhv import (
    HousingAppeal,
    HousingContract,
    HousingContractConsent,
    HousingPersonalAccount,
    HousingQueueCase,
    OrphanHousingRecord,
    YoungFamilyRecord,
)


PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")


def build_case_document_context(case: HousingQueueCase) -> dict[str, Any]:
    citizen = case.citizen
    members = list(case.household_members.all())
    household_lines = "\n".join(
        f"- {m.full_name} ({m.get_relation_display()})"
        + (f", доход {m.monthly_income} ₽" if m.monthly_income else "")
        for m in members
    )
    passport = ""
    if citizen.passport_series or citizen.passport_number:
        passport = f"{citizen.passport_series} {citizen.passport_number}".strip()
        if citizen.passport_issued_at:
            passport += f", выдан {citizen.passport_issued_at:%d.%m.%Y}"
        if citizen.passport_issued_by:
            passport += f", {citizen.passport_issued_by}"

    ctx = {
        "case_number": case.case_number,
        "citizen_name": citizen.full_name,
        "citizen_snils": citizen.snils,
        "citizen_address": citizen.reg_address,
        "passport": passport,
        "category": case.get_category_display(),
        "registered_at": case.registered_at.strftime("%d.%m.%Y"),
        "queue_position": str(case.queue_position or "—"),
        "household_size": str(case.household_size or len(members) or 1),
        "monthly_income": str(case.monthly_income or ""),
        "property_value": str(case.property_value or ""),
        "per_capita_income": str(case.per_capita_income or ""),
        "low_income_conclusion": case.low_income_conclusion or "",
        "household_list": household_lines or citizen.full_name,
        "today": timezone_now_str(),
    }

    if case.low_income_eligible is True:
        ctx["account_decision"] = "ПРИНЯТЬ на учёт в качестве нуждающегося"
        ctx["account_decision_title"] = "РЕШЕНИЕ о принятии на учёт"
    elif case.low_income_eligible is False:
        ctx["account_decision"] = "ОТКАЗАТЬ в постановке на учёт"
        ctx["account_decision_title"] = "РЕШЕНИЕ об отказе в постановке на учёт"
    else:
        ctx["account_decision"] = "Требуется расчёт малоимущих"
        ctx["account_decision_title"] = "ПРОЕКТ решения"

    if case.removal_reason:
        ctx["removal_reason"] = case.get_removal_reason_display()
        ctx["removed_at"] = case.removed_at.strftime("%d.%m.%Y") if case.removed_at else ""
    else:
        ctx["removal_reason"] = ""
        ctx["removed_at"] = ""

    yf = YoungFamilyRecord.objects.filter(case=case).first()
    if yf:
        ctx["spouse_name"] = yf.spouse_full_name
        ctx["marriage_date"] = yf.marriage_date.strftime("%d.%m.%Y") if yf.marriage_date else ""
        ctx["children_count"] = str(yf.children_count)
        ctx["young_family_program"] = yf.get_program_display()
        ctx["young_family_criteria"] = yf.criteria_notes or (
            "Соответствует критериям" if yf.meets_criteria else "Не соответствует"
        )
    else:
        ctx["spouse_name"] = ctx["marriage_date"] = ctx["children_count"] = ""
        ctx["young_family_program"] = ctx["young_family_criteria"] = ""

    orphan = OrphanHousingRecord.objects.filter(case=case).select_related("assigned_premise").first()
    if orphan:
        ctx["mintrud_decision_number"] = orphan.mintrud_decision_number
        ctx["mintrud_decision_date"] = (
            orphan.mintrud_decision_date.strftime("%d.%m.%Y") if orphan.mintrud_decision_date else ""
        )
        ctx["orphan_housing_status"] = orphan.get_housing_status_display()
        ctx["assigned_premise"] = str(orphan.assigned_premise) if orphan.assigned_premise else ""
    else:
        ctx["mintrud_decision_number"] = ctx["mintrud_decision_date"] = ""
        ctx["orphan_housing_status"] = ctx["assigned_premise"] = ""

    return ctx


def timezone_now_str() -> str:
    from django.utils import timezone

    return timezone.now().strftime("%d.%m.%Y")


def render_template_body(body: str, context: dict[str, Any]) -> str:
    def repl(match):
        key = match.group(1)
        return str(context.get(key, ""))

    return PLACEHOLDER_RE.sub(repl, body)


def get_uzhv_print_template(subsystem, code: str) -> PrintTemplate | None:
    return PrintTemplate.objects.filter(subsystem=subsystem, code=code, is_active=True).first()


def render_case_document(case: HousingQueueCase, template_code: str) -> tuple[str, str]:
    """Возвращает (title, rendered_text)."""
    tpl = get_uzhv_print_template(case.subsystem, template_code)
    if not tpl:
        raise KeyError(template_code)
    ctx = build_case_document_context(case)
    return tpl.name, render_template_body(tpl.body, ctx)


def build_appeal_document_context(appeal: HousingAppeal) -> dict[str, Any]:
    citizen = appeal.citizen
    return {
        "appeal_number": appeal.appeal_number,
        "received_at": appeal.received_at.strftime("%d.%m.%Y"),
        "due_date": appeal.due_date.strftime("%d.%m.%Y"),
        "citizen_name": citizen.full_name if citizen else "—",
        "citizen_address": citizen.reg_address if citizen else "",
        "subject": appeal.subject,
        "body": appeal.body or "",
        "answer_text": appeal.answer_text or "",
        "answered_at": appeal.answered_at.strftime("%d.%m.%Y") if appeal.answered_at else "",
        "conclusion_kind": appeal.get_conclusion_kind_display() if appeal.conclusion_kind else "",
        "outgoing_number": (
            appeal.outgoing_correspondence.reg_number if appeal.outgoing_correspondence_id else ""
        ),
        "housing_case_number": appeal.housing_case.case_number if appeal.housing_case_id else "",
        "today": timezone_now_str(),
    }


def render_appeal_document(appeal: HousingAppeal, template_code: str = "uzhv_appeal_response") -> tuple[str, str]:
    tpl = get_uzhv_print_template(appeal.subsystem, template_code)
    if not tpl:
        raise KeyError(template_code)
    ctx = build_appeal_document_context(appeal)
    return tpl.name, render_template_body(tpl.body, ctx)


CONSENT_TEMPLATE_CODES = {
    HousingContractConsent.ConsentType.SUBLET: "uzhv_consent_sublet",
    HousingContractConsent.ConsentType.MOVE_IN: "uzhv_consent_move_in",
    HousingContractConsent.ConsentType.EXCHANGE: "uzhv_consent_exchange",
    HousingContractConsent.ConsentType.TEMP_BAN: "uzhv_temp_residents_ban",
    HousingContractConsent.ConsentType.TERMINATION_OBLIGATION: "uzhv_termination_obligation",
    HousingContractConsent.ConsentType.EMERGENCY_AGREEMENT: "uzhv_emergency_agreement",
    HousingContractConsent.ConsentType.PRIVATIZATION: "uzhv_privatization_transfer",
    HousingContractConsent.ConsentType.PRIVATE_TO_MUNICIPAL: "uzhv_private_to_municipal",
}


def build_contract_document_context(
    contract: HousingContract, consent: HousingContractConsent | None = None
) -> dict[str, Any]:
    citizen = contract.citizen
    premise = str(contract.premise) if contract.premise else "—"
    ctx = {
        "contract_number": contract.contract_number,
        "contract_type": contract.get_contract_type_display(),
        "citizen_name": citizen.full_name,
        "citizen_address": citizen.reg_address,
        "premise_address": premise,
        "signed_at": contract.signed_at.strftime("%d.%m.%Y"),
        "today": timezone_now_str(),
        "consent_type": "",
        "decision": "",
        "subject": "",
        "document_number": "",
        "registered_at": "",
    }
    if consent:
        ctx.update(
            {
                "consent_type": consent.get_consent_type_display(),
                "decision": consent.get_decision_display(),
                "subject": consent.subject,
                "document_number": consent.document_number,
                "registered_at": consent.registered_at.strftime("%d.%m.%Y"),
            }
        )
    return ctx


def render_consent_document(consent: HousingContractConsent) -> tuple[str, str]:
    code = CONSENT_TEMPLATE_CODES.get(consent.consent_type)
    if not code:
        raise KeyError(consent.consent_type)
    contract = consent.contract
    tpl = get_uzhv_print_template(contract.subsystem, code)
    if not tpl:
        raise KeyError(code)
    ctx = build_contract_document_context(contract, consent)
    return tpl.name, render_template_body(tpl.body, ctx)


def render_personal_account_document(account: HousingPersonalAccount) -> tuple[str, str]:
    from delayu.services.uzhv_personal_account import build_account_extract_context

    tpl = get_uzhv_print_template(account.subsystem, "uzhv_personal_account_extract")
    if not tpl:
        raise KeyError("uzhv_personal_account_extract")
    ctx = build_account_extract_context(account)
    history_lines = "\n".join(
        f"- {h.changed_at:%d.%m.%Y %H:%M}: {h.description}"
        for h in account.history.select_related("changed_by").order_by("-changed_at")[:15]
    )
    ctx["history_list"] = history_lines or "—"
    return tpl.name, render_template_body(tpl.body, ctx)


def text_to_docx_bytes(title: str, paragraphs: str) -> bytes:
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)
        for block in paragraphs.split("\n\n"):
            for line in block.split("\n"):
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        pass
    # Fallback: minimal docx-like zip not worth it — return UTF-8 text wrapped
    return paragraphs.encode("utf-8")


UZHV_TEMPLATE_DEFAULTS = {
    "uzhv_low_income_conclusion": {
        "name": "Заключение о малоимущих",
        "body": """ЗАКЛЮЧЕНИЕ
о признании (непризнании) гражданина малоимущим

Дело № {{case_number}} от {{registered_at}}
Гражданин: {{citizen_name}}
СНИЛС: {{citizen_snils}}
Паспорт: {{passport}}
Адрес регистрации: {{citizen_address}}

Состав семьи:
{{household_list}}

Число членов семьи: {{household_size}}
Суммарный доход: {{monthly_income}} ₽
Среднедушевой доход: {{per_capita_income}} ₽
Стоимость имущества: {{property_value}} ₽

{{low_income_conclusion}}

Дата формирования: {{today}}""",
    },
    "uzhv_queue_certificate": {
        "name": "Справка о постановке на учёт",
        "body": """СПРАВКА

Подтверждается, что гражданин {{citizen_name}} состоит на учёте нуждающихся
в жилых помещениях под № {{case_number}} с {{registered_at}}.
Категория учёта: {{category}}.
Очерёдность: {{queue_position}}.

Дата выдачи: {{today}}""",
    },
    "uzhv_orphan_package_cover": {
        "name": "Сопроводительный лист пакета (дети-сироты)",
        "body": """СОПРОВОДИТЕЛЬНЫЙ ЛИСТ
электронного пакета документов

Дело № {{case_number}}
Гражданин: {{citizen_name}}
Категория: {{category}}
Дата постановки на учёт: {{registered_at}}

Пакет сформирован: {{today}}
Состав: manifest.json, summary.txt, межведомственные запросы, вложения (при наличии).

Ответственный исполнитель: ______________________""",
    },
    "uzhv_account_decision": {
        "name": "Решение о постановке на учёт / отказе",
        "body": """{{account_decision_title}}

Дело № {{case_number}}
Заявитель: {{citizen_name}}
Адрес: {{citizen_address}}
Дата постановки на учёт: {{registered_at}}

{{account_decision}}

Основание: {{low_income_conclusion}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_young_family_certificate": {
        "name": "Справка для списка молодых семей",
        "body": """СПРАВКА
для включения в список молодых семей

Дело № {{case_number}}
Заявитель: {{citizen_name}}
Супруг(а): {{spouse_name}}
Дата брака: {{marriage_date}}
Детей: {{children_count}}
Программа: {{young_family_program}}

Заключение по критериям:
{{young_family_criteria}}

Дата: {{today}}""",
    },
    "uzhv_orphan_resolution_draft": {
        "name": "Проект постановления (дети-сироты)",
        "body": """ПРОЕКТ ПОСТАНОВЛЕНИЯ
о реализации жилищных прав

Дело № {{case_number}}
Гражданин: {{citizen_name}}

Решение Минтруда КК № {{mintrud_decision_number}} от {{mintrud_decision_date}}
Статус: {{orphan_housing_status}}
Специализированное помещение: {{assigned_premise}}

ПРЕДЛАГАЕТСЯ:
утвердить предоставление жилого помещения / выплаты в соответствии с решением Минтруда.

Дата проекта: {{today}}
Исполнитель: ______________________""",
    },
    "uzhv_consent_sublet": {
        "name": "Согласие / отказ на поднайм",
        "body": """{{decision}} НА ПОДНАЙМ ЖИЛОГО ПОМЕЩЕНИЯ

Договор № {{contract_number}} от {{signed_at}}
Наниматель: {{citizen_name}}
Помещение: {{premise_address}}

Содержание: {{subject}}

Дата регистрации: {{registered_at}}
№ документа: {{document_number}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_consent_move_in": {
        "name": "Согласие / отказ на вселение",
        "body": """{{decision}} НА ВСЕЛЕНИЕ ЧЛЕНОВ СЕМЬИ

Договор № {{contract_number}} от {{signed_at}}
Наниматель: {{citizen_name}}
Помещение: {{premise_address}}

Вселяемые: {{subject}}

Дата регистрации: {{registered_at}}
№ документа: {{document_number}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_consent_exchange": {
        "name": "Согласие на обмен жилого помещения",
        "body": """СОГЛАСИЕ НА ОБМЕН ЖИЛОГО ПОМЕЩЕНИЯ

Договор № {{contract_number}} от {{signed_at}}
Наниматель: {{citizen_name}}
Помещение: {{premise_address}}

Условия обмена: {{subject}}

Дата регистрации: {{registered_at}}
№ документа: {{document_number}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_temp_residents_ban": {
        "name": "Запрет проживания временных жильцов",
        "body": """УВЕДОМЛЕНИЕ О ЗАПРЕТЕ ПРОЖИВАНИЯ ВРЕМЕННЫХ ЖИЛЬЦОВ

Договор № {{contract_number}} от {{signed_at}}
Наниматель: {{citizen_name}}
Помещение: {{premise_address}}

Основание: {{subject}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_termination_obligation": {
        "name": "Обязательство о расторжении договора",
        "body": """ОБЯЗАТЕЛЬСТВО О РАСТОРЖЕНИИ ДОГОВОРА И ОСВОБОЖДЕНИИ ЖИЛЬЯ

Договор № {{contract_number}} от {{signed_at}}
Наниматель: {{citizen_name}}
Помещение: {{premise_address}}

Условия: {{subject}}

Дата регистрации: {{registered_at}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_emergency_agreement": {
        "name": "Соглашение об изъятии (аварийный МКД)",
        "body": """СОГЛАШЕНИЕ ОБ ИЗЪЯТИИ НЕДВИЖИМОСТИ

Договор № {{contract_number}} от {{signed_at}}
Гражданин: {{citizen_name}}
Помещение: {{premise_address}}

Условия изъятия: {{subject}}

Дата регистрации: {{registered_at}}
№ документа: {{document_number}}

Дата: {{today}}
Подпись: ______________________""",
    },
    "uzhv_privatization_transfer": {
        "name": "Договор передачи жилья в собственность",
        "body": """РЕГИСТРАЦИЯ ПЕРЕДАЧИ ЖИЛОГО ПОМЕЩЕНИЯ В СОБСТВЕННОСТЬ

Договор найма № {{contract_number}} от {{signed_at}}
Гражданин: {{citizen_name}}
Помещение: {{premise_address}}

Примечание: {{subject}}

Дата регистрации: {{registered_at}}

Дата: {{today}}""",
    },
    "uzhv_private_to_municipal": {
        "name": "Безвозмездная передача в муниципальную собственность",
        "body": """РЕГИСТРАЦИЯ БЕЗВОЗМЕЗДНОЙ ПЕРЕДАЧИ В МУНИЦИПАЛЬНУЮ СОБСТВЕННОСТЬ

Объект: {{subject}}
Связанный договор: № {{contract_number}}

Дата регистрации: {{registered_at}}
№ документа: {{document_number}}

Дата: {{today}}""",
    },
    "uzhv_personal_account_extract": {
        "name": "Выписка из лицевого счёта",
        "body": """ВЫПИСКА ИЗ ЛИЦЕВОГО СЧЁТА
Лицевой счёт № {{account_number}} ({{account_status}})
Дата открытия: {{opened_at}}

Адрес помещения: {{premise_address}}
МКД: {{building_address}}
Жилая / общая площадь: {{living_area}} / {{total_area}} м²
Комнат: {{rooms}}

Наниматель (собственник): {{tenant_name}}
СНИЛС: {{tenant_snils}}
Адрес регистрации: {{tenant_address}}

Состав семьи:
{{members_list}}

Коммунальные услуги:
{{utility_services}}

История изменений:
{{history_list}}

Дата выдачи выписки: {{today}}
Подпись: ______________________""",
    },
    "uzhv_appeal_response": {
        "name": "Ответ на обращение гражданина",
        "body": """ОТВЕТ НА ОБРАЩЕНИЕ ГРАЖДАНИНА

Обращение № {{appeal_number}} от {{received_at}}
Заявитель: {{citizen_name}}
Адрес: {{citizen_address}}
Учётное дело: {{housing_case_number}}

Тема: {{subject}}

Содержание обращения:
{{body}}

Вид заключения: {{conclusion_kind}}

ТЕКСТ ОТВЕТА:
{{answer_text}}

Исходящий № {{outgoing_number}}
Дата ответа: {{answered_at}}

Дата: {{today}}
Подпись: ______________________""",
    },
}


def seed_uzhv_print_templates(subsystem) -> None:
    for code, meta in UZHV_TEMPLATE_DEFAULTS.items():
        PrintTemplate.objects.update_or_create(
            subsystem=subsystem,
            code=code,
            defaults={"name": meta["name"], "body": meta["body"], "is_active": True},
        )
